import io
import random
from PIL import Image
from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Function to hide image within Word document using semagram-based steganography
def hide_image_in_word(image_file, word_file):
    # Open the image file and read its contents
    with open(image_file, 'rb') as f:
        image_bytes = f.read()

    # Open the Word document file and create a Document object
    document = Document(word_file)

    # Get the text content of the document
    text_content = ''
    for paragraph in document.paragraphs:
        text_content += paragraph.text

    # Generate the semagrams for the image bytes
    semagram_size = 3
    semagrams = generate_semagrams(image_bytes, semagram_size)

    # Calculate the size of the image
    image_size = Image.open(io.BytesIO(image_bytes)).size

    # Calculate the maximum number of semagrams that can be hidden
    max_semagrams = len(text_content) // (semagram_size * 2)

    # If there are more semagrams than can be hidden, raise an exception
    if len(semagrams) > max_semagrams:
        raise ValueError('Too many semagrams to hide in text')

    # Insert the semagrams into the text content at random positions
    modified_content = insert_semagrams_in_text(semagrams, text_content)

    # Replace the original text content with the modified content
    for i, paragraph in enumerate(document.paragraphs):
        paragraph.text = modified_content[i]

    # Add an invisible image to the Word document with the same size as the original image
    invisible_image = Image.new('RGB', image_size, color='white')
    with io.BytesIO() as image_stream:
        invisible_image.save(image_stream, 'PNG')
        image_stream.seek(0)
        document.add_picture(image_stream, width=Cm(image_size[0]), height=Cm(image_size[1]))

    # Save the modified document
    document.save(word_file)

# Function to generate semagrams from bytes
def generate_semagrams(data, size=3):
    semagrams = []
    for i in range(0, len(data), size):
        semagrams.append(data[i:i+size])
    return semagrams

# Function to insert semagrams into text at random positions
def insert_semagrams_in_text(semagrams, text):
    # Convert the text to a list of characters
    characters = list(text)

    # Calculate the number of semagrams that can be hidden in the text
    max_semagrams = len(text) // (len(semagrams) * 2)

    # If there are more semagrams than can be hidden, raise an exception
    if len(semagrams) > max_semagrams:
        raise ValueError('Too many semagrams to hide in text')

    # Calculate the dimensions of the rectangle for hiding the image
    num_rows = int(len(semagrams) ** 0.5)
    num_cols = int(len(semagrams) / num_rows)
    if num_rows * num_cols < len(semagrams):
        num_cols += 1

    # Fill the rectangle with the semagrams
    for i, semagram in enumerate(semagrams):
        row = i // num_cols
        col = i % num_cols
        x = (col * 2) + 1
        y = (row * 2) + 1
        index = (y * (num_cols * 2)) +


# Example usage
if __name__ == '__main__':
    # Input file paths
    image_file = 'image.png'
    word_file = 'text.docx'

    # Hide the image within the Word document
    hide_image_in_word(image_file, word_file, semagram_size=15)
